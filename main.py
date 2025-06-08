import os, uuid, shutil, pathlib, asyncio, re, io
from fastapi import FastAPI, UploadFile, Form, BackgroundTasks, Request, Response
from fastapi.responses import HTMLResponse, StreamingResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
import openai

openai.api_key = os.getenv("OPENAI_API_KEY")

BASE_DIR = pathlib.Path(__file__).parent
TMP_DIR = BASE_DIR / "tmp"
TMP_DIR.mkdir(exist_ok=True)

app = FastAPI()
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))

TASKS = {}  # task_id -> {"progress":float,"ats":float,"file":Path}

# ----- basic pages --------------------------------------------------------
@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

# ----- helpers ------------------------------------------------------------
def extract_bullets(doc_path: pathlib.Path):
    doc = Document(str(doc_path))
    bullets = []
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("List"):
            bullets.append(p.text.strip())
    return bullets

async def rewrite_bullet(text:str, jd:str):
    if not openai.api_key:
        return text + " - tailored"
    prompt = f"You are a resume copy-editor. Rewrite the line so it matches this JD snippet:\n---\n{jd}\n---\nOriginal: {text}\nRewritten:"
    resp = await openai.ChatCompletion.acreate(
        model="gpt-4o-mini",
        temperature=0,
        messages=[{"role":"user","content":prompt}]
    )
    return resp.choices[0].message.content.strip()

def compute_ats(bullets, jd):
    jd_keywords = {w.lower() for w in re.findall(r"[A-Za-z]+", jd)}
    hit = sum(any(k in b.lower() for k in jd_keywords) for b in bullets)
    return round(10 * hit/len(bullets), 1) if bullets else 0

def save_docx(template_path:pathlib.Path, new_bullets:list[str], out_path:pathlib.Path):
    doc = Document(str(template_path))
    it = iter(new_bullets)
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("List"):
            p.text = next(it)
    doc.save(str(out_path))

# ----- background task ----------------------------------------------------
async def do_rewrite(task_id:str, resume_path:pathlib.Path, jd:str):
    bullets = extract_bullets(resume_path)
    rewritten = []
    total = len(bullets)
    for i, b in enumerate(bullets, 1):
        rewritten.append(await rewrite_bullet(b, jd))
        TASKS[task_id]["progress"] = round(100*i/total, 1)
    ats = compute_ats(rewritten, jd)
    out_file = TMP_DIR / f"{task_id}.docx"
    save_docx(resume_path, rewritten, out_file)
    TASKS[task_id].update({"ats": ats, "file": out_file})

# ----- API ----------------------------------------------------------------
@app.post("/generate")
async def generate(background_tasks: BackgroundTasks,
                   file: UploadFile,
                   job_description: str = Form(...)):
    task_id = uuid.uuid4().hex
    tmp_resume = TMP_DIR / f"{task_id}_{file.filename}"
    with tmp_resume.open("wb") as f:
        shutil.copyfileobj(file.file, f)
    TASKS[task_id] = {"progress": 0.0, "ats": None, "file": None}
    background_tasks.add_task(do_rewrite, task_id, tmp_resume, job_description[:4000])
    return {"task_id": task_id}

@app.get("/status/{task_id}")
def status(task_id:str):
    task = TASKS.get(task_id)
    if not task:
        return JSONResponse(status_code=404, content={"error":"unknown id"})
    return {"progress": task["progress"], "ats": task["ats"]}

@app.get("/download/{task_id}")
async def download(task_id:str):
    task = TASKS.get(task_id)
    if not task or not task["file"]:
        return JSONResponse(status_code=404, content={"error":"not ready"})
    def file_stream():
        with open(task["file"], "rb") as f:
            yield from iter(lambda: f.read(8192), b"")
    headers = {"Content-Disposition": f'attachment; filename="tailored_{task_id}.docx"'}
    return StreamingResponse(file_stream(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers=headers)
